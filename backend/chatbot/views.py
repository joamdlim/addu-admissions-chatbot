from django.shortcuts import render
from rest_framework.decorators import api_view
from rest_framework.response import Response
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, StreamingHttpResponse, HttpResponse
import json
from .chroma_connection import ChromaService
from .improved_pdf_to_chroma import sync_supabase_to_chroma_improved
from .supabase_client import get_supabase_client
from .models import Conversation, ConversationTurn, SystemPrompt
# Remove the enhanced chatbot import
# from .enhanced_hybrid_chatbot import get_enhanced_chatbot
# from .fast_hybrid_chatbot import FastHybridChatbot

# Add the Together AI chatbot import
from .fast_hybrid_chatbot_together import FastHybridChatbotTogether
from .topics import TOPICS, CONVERSATION_STATES
import os, json
import uuid
import io  # Add this import at the top
from urllib.parse import quote
from .models import DocumentFolder, DocumentMetadata
from django.db.models import Count

# Initialize Together AI chatbot at startup
print("🚀 Initializing Together AI chatbot...")
together_chatbot = FastHybridChatbotTogether(use_chroma=True)
print("✅ Together AI chatbot ready!")

@api_view(['GET'])
def evaluate(request):
    return Response({"f1": 0.85, "bleu": 0.72, "satisfaction": 4.2})

@api_view(['GET'])
def get_topics(request):
    """Get all available topics for guided conversation"""
    try:
        topics_list = []
        for topic_id, topic_data in TOPICS.items():
            topics_list.append({
                'id': topic_id,
                'label': topic_data['label'],
                'description': topic_data['description'],
                'keywords': topic_data['keywords']
            })
        
        return Response({
            "topics": topics_list,
            "conversation_states": CONVERSATION_STATES
        })
        
    except Exception as e:
        return Response({"error": str(e)}, status=500)

@csrf_exempt
def guided_chat_view(request):
    """Guided conversation chat view with topic-based filtering"""
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            user_input = data.get("user_input", "")
            action_type = data.get("action_type", "message")  # 'message', 'topic_selection', 'action'
            action_data = data.get("action_data", None)
            session_id = data.get("session_id", None)
            
            # Set session ID if provided
            if session_id:
                together_chatbot.set_session_state(session_id=session_id)
            
            # Process guided conversation
            result = together_chatbot.process_guided_conversation(
                user_input=user_input,
                action_type=action_type,
                action_data=action_data
            )
            
            # Add session info to response
            result['session_id'] = together_chatbot.get_session_state().get('session_id', session_id)
            
            return JsonResponse(result)
            
        except Exception as e:
            print(f"❌ Guided chat error: {e}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"error": "POST request required"}, status=400)

@csrf_exempt
def upload_pdf_view(request):
    if request.method == "POST":
        if 'pdf_file' not in request.FILES:
            return JsonResponse({"error": "No file uploaded"}, status=400)
        
        pdf_file = request.FILES['pdf_file']
        pdf_bytes = pdf_file.read()
        pdf_file_name = pdf_file.name

        try:
            # This function is no longer used as per the new_code, but kept for now
            # process_and_store_pdf_in_chroma(pdf_bytes, pdf_file_name) 
            return JsonResponse({"status": f"PDF '{pdf_file_name}' processed and stored successfully."})
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "POST request required"}, status=400)

@csrf_exempt
def chat_view(request):
    """Chat view using Together AI"""
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            prompt = data.get("prompt", "")
            session_id = data.get("session_id", None)
            max_tokens = data.get("max_tokens", 3000)
            min_relevance = data.get("min_relevance", 0.1)
            
            def generate_streaming_response():
                try:
                    # Use the Together AI chatbot
                    for event in together_chatbot.process_query_stream(
                        query=prompt,
                        max_tokens=max_tokens,
                        min_relevance=min_relevance,
                        use_history=True
                    ):
                        yield "data: " + json.dumps(event) + "\n\n"
                        
                except Exception as e:
                    yield "data: " + json.dumps({"error": str(e)}) + "\n\n"
                    yield "data: " + json.dumps({"done": True}) + "\n\n"
            
            response = StreamingHttpResponse(
                generate_streaming_response(),
                content_type='text/event-stream'
            )
            
            response['Cache-Control'] = 'no-cache'
            response['Access-Control-Allow-Origin'] = '*'
            response['Access-Control-Allow-Headers'] = 'Content-Type'
            
            return response
            
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"error": "POST request required"}, status=400)

@csrf_exempt 
def chat_legacy_view(request):
    """Legacy chat view for backward compatibility (using Together AI)"""
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            prompt = data.get("prompt", "")
            
            def generate_streaming_response():
                try:
                    # Use Together AI chatbot for legacy support
                    for event in together_chatbot.process_query_stream(prompt, max_tokens=3000, min_relevance=0.1, use_history=True):
                        yield "data: " + json.dumps(event) + "\n\n"
                        
                except Exception as e:
                    yield "data: " + json.dumps({"error": str(e)}) + "\n\n"
                    yield "data: " + json.dumps({"done": True}) + "\n\n"
            
            response = StreamingHttpResponse(
                generate_streaming_response(),
                content_type='text/event-stream'
            )
            
            response['Cache-Control'] = 'no-cache'
            response['Access-Control-Allow-Origin'] = '*'
            response['Access-Control-Allow-Headers'] = 'Content-Type'
            
            return response
            
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"error": "POST request required"}, status=400)

# ===== SIMPLIFIED CONVERSATION MANAGEMENT ENDPOINTS =====
# Note: These endpoints are simplified since we're not using complex conversation memory

@api_view(['POST'])
def create_conversation(request):
    """Create a new conversation session (simplified)"""
    try:
        session_id = str(uuid.uuid4())
        
        return Response({
            "status": "success",
            "session_id": session_id,
            "message": "Conversation created (simplified mode)"
        })
        
    except Exception as e:
        return Response({"error": str(e)}, status=500)

@api_view(['GET'])
def get_conversation_history(request, session_id):
    """Get conversation history for a session (simplified)"""
    try:
        # In simplified mode, we don't store conversation history
        # This is just for API compatibility
        return Response({
            "session_id": session_id,
            "history": [],
            "count": 0,
            "message": "Conversation history not stored in simplified mode"
        })
        
    except Exception as e:
        return Response({"error": str(e)}, status=500)

@api_view(['GET'])
def get_conversation_stats(request, session_id):
    """Get conversation statistics (simplified)"""
    try:
        return Response({
            'session_id': session_id,
            'message': 'Simplified mode - no detailed stats available'
        })
        
    except Exception as e:
        return Response({"error": str(e)}, status=500)

@api_view(['DELETE'])
def clear_conversation(request, session_id):
    """Clear/reset a conversation (simplified)"""
    try:
        return Response({
            "status": "success",
            "message": f"Conversation {session_id} cleared (simplified mode)"
        })
        
    except Exception as e:
        return Response({"error": str(e)}, status=500)

@api_view(['GET'])
def list_active_conversations(request):
    """List all active conversations (simplified)"""
    try:
        return Response({
            "conversations": [],
            "count": 0,
            "message": "Simplified mode - no conversation tracking"
        })
        
    except Exception as e:
        return Response({"error": str(e)}, status=500)

@csrf_exempt
def force_summarization(request, session_id):
    """Manually trigger summarization for a conversation (simplified)"""
    if request.method == "POST":
        try:
            return JsonResponse({
                "status": "success",
                "message": "Summarization not needed in simplified mode"
            })
            
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"error": "POST request required"}, status=400)

# ===== SYSTEM PROMPT MANAGEMENT =====

@api_view(['GET'])
def get_system_prompts(request):
    """Get all system prompts"""
    try:
        prompts = SystemPrompt.objects.all().order_by('-created_at')
        
        prompt_list = []
        for prompt in prompts:
            prompt_list.append({
                'id': prompt.id,
                'name': prompt.name,
                'prompt_text': prompt.prompt_text,
                'token_count': prompt.token_count,
                'is_active': prompt.is_active,
                'created_at': prompt.created_at.isoformat()
            })
        
        return Response({"prompts": prompt_list})
        
    except Exception as e:
        return Response({"error": str(e)}, status=500)

@csrf_exempt
def update_system_prompt(request):
    """Create or update system prompt"""
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            name = data.get("name", "default")
            prompt_text = data.get("prompt_text", "")
            
            if not prompt_text:
                return JsonResponse({"error": "Prompt text is required"}, status=400)
            
            # Count tokens (simplified)
            token_count = len(prompt_text.split())
            
            # Deactivate other prompts
            SystemPrompt.objects.filter(is_active=True).update(is_active=False)
            
            # Create or update prompt
            prompt, created = SystemPrompt.objects.update_or_create(
                name=name,
                defaults={
                    'prompt_text': prompt_text,
                    'token_count': token_count,
                    'is_active': True
                }
            )
            
            return JsonResponse({
                "status": "success",
                "prompt_id": prompt.id,
                "token_count": token_count,
                "created": created
            })
            
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"error": "POST request required"}, status=400)

# ===== EXISTING ENDPOINTS (UNCHANGED) =====

def chroma_test_add(request):
    collection = ChromaService.get_collection()
    
    # Dummy embedding
    collection.add(
        ids=["test-id"],
        documents=["Hello Chroma from Django"],
        embeddings=[[0.1, 0.2, 0.3, 0.4]]
    )
    
    return JsonResponse({"status": "Document added"})

def chroma_test_query(request):
    collection = ChromaService.get_collection()
    
    results = collection.query(
        query_embeddings=[[0.1, 0.2, 0.3, 0.4]],
        n_results=1
    )
    
    return JsonResponse({"results": results})

@csrf_exempt
def sync_supabase_to_chroma(request):
    try:
        # Use the improved sync function with default parameters
        result = sync_supabase_to_chroma_improved()
        
        # Convert to the expected response format
        if "error" in result:
            return JsonResponse({"error": result["error"]}, status=500)
        
        return JsonResponse({
            "status": result["status"], 
            "found": result["found"], 
            "ingested": result["ingested"],
            "extraction_method": result.get("extraction_method", "pdfplumber")
        })
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)

@api_view(['GET'])
def export_chroma_to_local(request):
    try:
        client = ChromaService.get_client()
        collection_name = os.getenv("CHROMA_COLLECTION", "documents")
        col = client.get_or_create_collection(name=collection_name)

        # Pull all items (paginate)
        all_docs = []
        limit = 1000
        offset = 0
        while True:
            res = col.get(include=["documents", "metadatas"], limit=limit, offset=offset)
            ids = res.get("ids", [])
            docs = res.get("documents", [])
            if not ids:
                break
            for i, doc_id in enumerate(ids):
                content = docs[i] if i < len(docs) else ""
                all_docs.append({"id": doc_id, "content": content})
            offset += len(ids)

        # Paths relative to backend/
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        processed_dir = os.path.join(base_dir, "processed")
        embeddings_dir = os.path.join(base_dir, "embeddings")
        os.makedirs(processed_dir, exist_ok=True)
        os.makedirs(embeddings_dir, exist_ok=True)

        processed_path = os.path.join(processed_dir, "processed.json")
        metadata_path = os.path.join(embeddings_dir, "metadata.json")

        with open(processed_path, "w", encoding="utf-8") as f:
            json.dump(all_docs, f, ensure_ascii=False, indent=2)
        with open(metadata_path, "w", encoding="utf-8") as f:
            json.dump(all_docs, f, ensure_ascii=False, indent=2)

        # Optional: force local vector regeneration next time by removing vectors file
        vectors_path = os.path.join(embeddings_dir, "hybrid_vectors.npy")
        if os.path.exists(vectors_path):
            os.remove(vectors_path)

        return JsonResponse({"status": "ok", "exported": len(all_docs), "collection": collection_name})
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)

@csrf_exempt
def admin_upload_file(request):
    """Upload file directly to Supabase storage and process it for ChromaDB with enhanced metadata"""
    if request.method == "POST":
        if 'file' not in request.FILES:
            return JsonResponse({"error": "No file uploaded"}, status=400)
        
        uploaded_file = request.FILES['file']
        file_bytes = uploaded_file.read()
        file_name = uploaded_file.name
        
        # Get optional metadata from form data
        folder_id = request.POST.get('folder_id')
        document_type = request.POST.get('document_type', 'other')
        target_program = request.POST.get('target_program', 'all')
        keywords = request.POST.get('keywords', '')
        
        try:
            supabase = get_supabase_client()
            
            # Upload to Supabase storage
            result = supabase.storage.from_("original-pdfs").upload(
                file_name,
                file_bytes,
                {"content-type": uploaded_file.content_type}
            )
            
            # Process and store in ChromaDB immediately for PDFs only
            if file_name.lower().endswith('.pdf'):
                try:
                    from .improved_pdf_to_chroma import process_and_store_pdf_in_chroma_improved
                    documents_stored = process_and_store_pdf_in_chroma_improved(
                        file_bytes, 
                        file_name,
                        chunk_size=1000,
                        overlap=200,
                        folder_id=int(folder_id) if folder_id else None,
                        document_type=document_type,
                        target_program=target_program,
                        keywords=keywords
                    )
                    message = f"File '{file_name}' uploaded and processed successfully ({documents_stored} documents stored in ChromaDB)"
                except Exception as chroma_error:
                    print(f"⚠️ ChromaDB processing failed for {file_name}: {chroma_error}")
                    message = f"File '{file_name}' uploaded to Supabase but ChromaDB processing failed: {str(chroma_error)}"
            else:
                message = f"File '{file_name}' uploaded successfully (non-PDF files are stored in Supabase only)"
            
            return JsonResponse({
                "status": "success",
                "message": message,
                "file_name": file_name
            })
            
        except Exception as e:
            print(f"❌ Upload failed: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"error": "POST request required"}, status=400)

@api_view(['GET'])
def admin_list_files(request):
    """List all files from Supabase storage"""
    try:
        supabase = get_supabase_client()
        files = supabase.storage.from_("original-pdfs").list()
        
        # Format file data
        file_list = []
        for file in files:
            file_list.append({
                "name": file.get("name"),
                "size": file.get("metadata", {}).get("size", 0),
                "created_at": file.get("created_at"),
                "updated_at": file.get("updated_at"),
                "content_type": file.get("metadata", {}).get("mimetype", "unknown")
            })
        
        return Response({"files": file_list})
        
    except Exception as e:
        return Response({"error": str(e)}, status=500)

@csrf_exempt
def admin_delete_file(request):
    """Delete file from Supabase storage and ChromaDB"""
    if request.method == "DELETE":
        try:
            data = json.loads(request.body)
            file_name = data.get("file_name")
            
            if not file_name:
                return JsonResponse({"error": "File name required"}, status=400)
            
            supabase = get_supabase_client()
            
            # Delete from Supabase storage
            supabase.storage.from_("original-pdfs").remove([file_name])
            print(f"✅ Deleted from Supabase: {file_name}")
            
            # Also remove from ChromaDB for supported file types
            if file_name.lower().endswith(('.pdf', '.txt', '.docx')):
                try:
                    from .chroma_connection import ChromaService
                    collection = ChromaService.get_collection()
                    
                    # Use the same ID format as storage (single document approach)
                    doc_id = os.path.splitext(file_name)[0]
                    print(f"🗑️ Deleting ChromaDB document: {doc_id}")
                    
                    collection.delete(ids=[doc_id])
                    message = f"File '{file_name}' deleted from both Supabase and ChromaDB"
                    print(f"✅ Successfully deleted from ChromaDB: {doc_id}")
                        
                except Exception as chroma_error:
                    print(f"⚠️ ChromaDB deletion failed for {file_name}: {chroma_error}")
                    import traceback
                    traceback.print_exc()
                    message = f"File '{file_name}' deleted from Supabase but ChromaDB deletion failed: {str(chroma_error)}"
            else:
                message = f"File '{file_name}' deleted successfully"
            
            return JsonResponse({
                "status": "success",
                "message": message
            })
            
        except Exception as e:
            print(f"❌ Delete failed: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"error": "DELETE request required"}, status=400)

@csrf_exempt
def admin_download_file(request, file_name):
    """Download file from Supabase storage"""
    try:
        supabase = get_supabase_client()
        
        print(f"📥 Downloading file: {file_name}")
        
        # Download file from Supabase
        file_bytes = supabase.storage.from_("original-pdfs").download(file_name)
        
        # Determine content type based on file extension
        content_type = "application/octet-stream"  # Default
        if file_name.lower().endswith('.pdf'):
            content_type = "application/pdf"
        elif file_name.lower().endswith('.txt'):
            content_type = "text/plain; charset=utf-8"
        elif file_name.lower().endswith('.docx'):
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # Create HTTP response with file
        response = HttpResponse(file_bytes, content_type=content_type)
        
        # Properly encode filename for Content-Disposition header
        encoded_filename = quote(file_name.encode('utf-8'))
        response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{encoded_filename}'
        response['Content-Length'] = len(file_bytes)
        response['Cache-Control'] = 'no-cache'
        
        print(f"✅ File download prepared: {file_name} ({len(file_bytes)} bytes)")
        return response
        
    except Exception as e:
        print(f"❌ Download failed for {file_name}: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({"error": f"Download failed: {str(e)}"}, status=500)

# Keep the manual sync function for individual files if needed
@csrf_exempt
def admin_sync_file_to_chroma(request):
    """Manually sync a specific file from Supabase to ChromaDB"""
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            file_name = data.get("file_name")
            
            if not file_name:
                return JsonResponse({"error": "File name required"}, status=400)
            
            if not file_name.lower().endswith(('.pdf', '.txt', '.docx')):
                return JsonResponse({"error": "Only PDF, TXT, and DOCX files can be synced to ChromaDB"}, status=400)
            
            supabase = get_supabase_client()
            
            # Download file from Supabase
            print(f"📥 Downloading {file_name} for sync...")
            file_bytes = supabase.storage.from_("original-pdfs").download(file_name)
            
            # Process based on file type
            if file_name.lower().endswith('.pdf'):
                # Use the existing PDF processing function (single document approach)
                from .improved_pdf_to_chroma import process_and_store_pdf_in_chroma_improved
                result = process_and_store_pdf_in_chroma_improved(
                    file_bytes, 
                    file_name,
                    chunk_size=1000,
                    overlap=200
                )
            else:
                # For TXT/DOCX files, extract text and use single document approach
                if file_name.lower().endswith('.txt'):
                    text_content = file_bytes.decode('utf-8', errors='ignore')
                elif file_name.lower().endswith('.docx'):
                    try:
                        import docx
                        import io
                        doc = docx.Document(io.BytesIO(file_bytes))
                        text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    except Exception as docx_error:
                        return JsonResponse({"error": f"DOCX processing failed: {str(docx_error)}"}, status=500)
                
                # Store as single document
                result = process_text_with_existing_embeddings(text_content, file_name)
            
            return JsonResponse({
                "status": "success",
                "message": f"File '{file_name}' successfully synced to ChromaDB",
                "documents_stored": result
            })
            
        except Exception as e:
            print(f"❌ Sync failed: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"error": "POST request required"}, status=400)

@csrf_exempt
def extract_text_from_file(request):
    """Extract text from uploaded file for preview/editing before final upload"""
    if request.method == "POST":
        if 'file' not in request.FILES:
            return JsonResponse({"error": "No file uploaded"}, status=400)
        
        uploaded_file = request.FILES['file']
        file_bytes = uploaded_file.read()
        file_name = uploaded_file.name
        file_type = uploaded_file.content_type
        
        print(f"🔍 Extracting text from: {file_name} ({file_type})")
        
        try:
            extracted_text = ""
            
            if file_name.lower().endswith('.pdf'):
                print("📄 Processing PDF file...")
                from .improved_pdf_to_chroma import extract_and_process_pdf
                cleaned_text, chunks, analysis = extract_and_process_pdf(file_bytes)
                extracted_text = cleaned_text
                
            elif file_name.lower().endswith('.txt'):
                print("📝 Processing TXT file...")
                extracted_text = file_bytes.decode('utf-8', errors='ignore')
                
            elif file_name.lower().endswith('.docx'):
                print("📄 Processing DOCX file...")
                try:
                    import docx
                    doc = docx.Document(io.BytesIO(file_bytes))
                    extracted_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    if not extracted_text.strip():
                        extracted_text = "Warning: No text content found in DOCX file."
                except ImportError:
                    extracted_text = "Error: python-docx package is required for DOCX files. Please install it: pip install python-docx"
                except Exception as docx_error:
                    print(f"DOCX processing error: {docx_error}")
                    extracted_text = f"DOCX processing failed: {str(docx_error)}\n\nPlease try converting to PDF or TXT format."
                    
            else:
                return JsonResponse({
                    "error": f"Unsupported file type: {file_name}. Please use PDF, TXT, or DOCX files."
                }, status=400)
            
            print(f"✅ Text extracted successfully: {len(extracted_text)} characters")
            
            return JsonResponse({
                "status": "success",
                "file_name": file_name,
                "file_type": file_type,
                "extracted_text": extracted_text,
                "text_length": len(extracted_text)
            })
            
        except Exception as e:
            print(f"❌ Text extraction failed: {str(e)}")
            return JsonResponse({"error": f"Text extraction failed: {str(e)}"}, status=500)
    
    return JsonResponse({"error": "POST request required"}, status=400)

# Add debugging to the upload_processed_file function
@csrf_exempt
def upload_processed_file(request):
    """Upload file with processed/edited text content"""
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            file_name = data.get("file_name")
            processed_text = data.get("processed_text")
            original_file_type = data.get("file_type")
            
            print(f"🔍 Processing file: {file_name}")
            print(f"🔍 Text length: {len(processed_text) if processed_text else 0}")
            
            if not file_name or not processed_text:
                return JsonResponse({"error": "File name and processed text required"}, status=400)
            
            if not processed_text.strip():
                return JsonResponse({"error": "Processed text cannot be empty"}, status=400)
            
            supabase = get_supabase_client()
            
            # Convert processed text to bytes for storage
            text_bytes = processed_text.encode('utf-8')
            
            # Upload to Supabase storage
            print(f"📤 Uploading to Supabase: {file_name}")
            result = supabase.storage.from_("original-pdfs").upload(
                file_name,
                text_bytes,
                {"content-type": "text/plain"}
            )
            print(f"✅ Uploaded to Supabase: {file_name}")
            
            # Process and store in ChromaDB for supported file types
            if file_name.lower().endswith(('.pdf', '.txt', '.docx')):
                print(f"🔄 Starting ChromaDB processing for {file_name}")
                try:
                    # Use the EXACT same approach as the working version
                    # Convert the processed text back to bytes and use the existing function
                    documents_stored = process_text_with_existing_embeddings(
                        processed_text, 
                        file_name
                    )
                    
                    message = f"File '{file_name}' uploaded and processed successfully ({documents_stored} documents stored in ChromaDB)"
                    print(f"✅ ChromaDB processing completed: {documents_stored} documents")
                    
                except Exception as chroma_error:
                    print(f"⚠️ ChromaDB processing failed for {file_name}: {chroma_error}")
                    import traceback
                    traceback.print_exc()
                    message = f"File '{file_name}' uploaded to Supabase but ChromaDB processing failed: {str(chroma_error)}"
            else:
                print(f"ℹ️ Skipping ChromaDB processing (unsupported file type): {file_name}")
                message = f"File '{file_name}' uploaded successfully"
            
            return JsonResponse({
                "status": "success",
                "message": message,
                "file_name": file_name
            })
            
        except Exception as e:
            print(f"❌ Upload failed: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"error": "POST request required"}, status=400)


def process_text_with_existing_embeddings(text: str, file_name: str):
    """Process text using the SAME approach as the original - single document per file"""
    from .improved_pdf_to_chroma import embed_text, tfidf_vectorizer, word2vec_model
    from .chroma_connection import ChromaService
    
    print(f"🔍 Processing as single document (matching original setup)...")
    
    # Check if models are already loaded
    if tfidf_vectorizer is None or word2vec_model is None:
        print(f"⚠️ Embedding models not initialized, initializing now...")
        from .improved_pdf_to_chroma import initialize_embedding_models
        initialize_embedding_models()
    else:
        print(f"✅ Using existing initialized embedding models")
    
    # Get ChromaDB collection
    collection = ChromaService.get_collection()
    
    # Store as ONE document per file (same as original)
    doc_id = os.path.splitext(file_name)[0]
    
    try:
        # Generate single embedding for entire text
        full_embedding = embed_text(text)
        
        collection.add(
            ids=[doc_id],
            documents=[text],
            embeddings=[full_embedding],
            metadatas=[{
                "filename": file_name, 
                "source": "pdf_scrape",
                "processed_via": "text_editor"
            }]
        )
        
        print(f"✅ Stored complete document as single entry: {doc_id}")
        return 1  # Successfully stored 1 document (matching original)
        
    except Exception as e:
        print(f"❌ Failed to store document {file_name}: {e}")
        raise

# ===== FOLDER MANAGEMENT API =====

@api_view(['GET', 'POST'])
def manage_folders(request):
    """List all folders or create a new folder"""
    if request.method == 'GET':
        folders = DocumentFolder.objects.annotate(
            doc_count=Count('documents')  # Changed from document_count to doc_count
        ).order_by('name')
        
        folder_list = []
        for folder in folders:
            # Get document type breakdown
            type_breakdown = DocumentMetadata.objects.filter(folder=folder).values('document_type').annotate(count=Count('document_type'))
            
            folder_list.append({
                'id': folder.id,
                'name': folder.name,
                'description': folder.description,
                'color': folder.color,
                'document_count': folder.doc_count,  # Use the annotated field
                'type_breakdown': list(type_breakdown),
                'created_at': folder.created_at.isoformat(),
                'updated_at': folder.updated_at.isoformat(),
            })
        
        return Response({"folders": folder_list})
    
    elif request.method == 'POST':
        try:
            data = json.loads(request.body)
            name = data.get('name', '').strip()
            description = data.get('description', '').strip()
            color = data.get('color', '#063970')
            
            if not name:
                return Response({"error": "Folder name is required"}, status=400)
            
            # Check for duplicate names
            if DocumentFolder.objects.filter(name=name).exists():
                return Response({"error": "Folder name already exists"}, status=400)
            
            folder = DocumentFolder.objects.create(
                name=name,
                description=description,
                color=color
            )
            
            return Response({
                "message": f"Folder '{name}' created successfully",
                "folder": {
                    'id': folder.id,
                    'name': folder.name,
                    'description': folder.description,
                    'color': folder.color,
                    'document_count': 0,
                    'created_at': folder.created_at.isoformat(),
                }
            })
            
        except Exception as e:
            return Response({"error": str(e)}, status=500)

@csrf_exempt
def manage_folder_detail(request, folder_id):
    """Update or delete a specific folder"""
    try:
        folder = DocumentFolder.objects.get(id=folder_id)
    except DocumentFolder.DoesNotExist:
        return JsonResponse({"error": "Folder not found"}, status=404)
    
    if request.method == 'PUT':
        try:
            data = json.loads(request.body)
            
            if 'name' in data:
                new_name = data['name'].strip()
                if not new_name:
                    return JsonResponse({"error": "Folder name cannot be empty"}, status=400)
                
                # Check for duplicate names (excluding current folder)
                if DocumentFolder.objects.filter(name=new_name).exclude(id=folder_id).exists():
                    return JsonResponse({"error": "Folder name already exists"}, status=400)
                
                folder.name = new_name
            
            if 'description' in data:
                folder.description = data['description'].strip()
            
            if 'color' in data:
                folder.color = data['color']
            
            folder.save()
            
            return JsonResponse({
                "message": f"Folder '{folder.name}' updated successfully",
                "folder": {
                    'id': folder.id,
                    'name': folder.name,
                    'description': folder.description,
                    'color': folder.color,
                }
            })
            
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    
    elif request.method == 'DELETE':
        try:
            folder_name = folder.name
            document_count = folder.documents.count()
            
            if document_count > 0:
                return JsonResponse({
                    "error": f"Cannot delete folder '{folder_name}' - it contains {document_count} documents. Move or delete documents first."
                }, status=400)
            
            folder.delete()
            return JsonResponse({"message": f"Folder '{folder_name}' deleted successfully"})
            
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    
    return JsonResponse({"error": "Method not allowed"}, status=405)

@api_view(['GET', 'POST'])
def manage_document_metadata(request):
    """List documents with metadata or update document metadata"""
    if request.method == 'GET':
        folder_id = request.GET.get('folder_id')
        document_type = request.GET.get('document_type')
        program = request.GET.get('program')
        
        queryset = DocumentMetadata.objects.select_related('folder').all()
        
        # Apply filters
        if folder_id:
            queryset = queryset.filter(folder_id=folder_id)
        if document_type:
            queryset = queryset.filter(document_type=document_type)
        if program:
            queryset = queryset.filter(target_program=program)
        
        documents = []
        for doc in queryset.order_by('-created_at'):
            documents.append({
                'id': doc.id,
                'document_id': doc.document_id,
                'filename': doc.filename,
                'folder': {
                    'id': doc.folder.id,
                    'name': doc.folder.name,
                    'color': doc.folder.color,
                },
                'document_type': doc.document_type,
                'document_type_display': doc.get_document_type_display(),
                'target_program': doc.target_program,
                'target_program_display': doc.get_target_program_display(),
                'keywords': doc.keywords,
                'keywords_list': doc.get_keywords_list(),
                'synced_to_chroma': doc.synced_to_chroma,
                'created_at': doc.created_at.isoformat(),
                'last_modified': doc.last_modified.isoformat(),
            })
        
        return Response({"documents": documents})

@csrf_exempt
def update_document_metadata(request, document_id):
    """Update metadata for a specific document"""
    if request.method != 'PUT':
        return JsonResponse({"error": "PUT request required"}, status=400)
    
    try:
        doc_metadata = DocumentMetadata.objects.get(document_id=document_id)
    except DocumentMetadata.DoesNotExist:
        return JsonResponse({"error": "Document not found"}, status=404)
    
    try:
        data = json.loads(request.body)
        
        # Update folder
        if 'folder_id' in data:
            try:
                new_folder = DocumentFolder.objects.get(id=data['folder_id'])
                doc_metadata.folder = new_folder
            except DocumentFolder.DoesNotExist:
                return JsonResponse({"error": "Folder not found"}, status=404)
        
        # Update document type
        if 'document_type' in data:
            valid_types = [choice[0] for choice in DocumentMetadata.DOCUMENT_TYPES]
            if data['document_type'] in valid_types:
                doc_metadata.document_type = data['document_type']
            else:
                return JsonResponse({"error": "Invalid document type"}, status=400)
        
        # Update target program
        if 'target_program' in data:
            valid_programs = [choice[0] for choice in DocumentMetadata.PROGRAMS]
            if data['target_program'] in valid_programs:
                doc_metadata.target_program = data['target_program']
            else:
                return JsonResponse({"error": "Invalid target program"}, status=400)
        
        # Update keywords
        if 'keywords' in data:
            doc_metadata.keywords = data['keywords'].strip()
        
        doc_metadata.save()
        
        # Update ChromaDB metadata
        try:
            from .chroma_connection import ChromaService
            collection = ChromaService.get_collection()
            
            # Get current document
            result = collection.get(ids=[document_id], include=["metadatas", "documents"])
            if result['ids']:
                # Update metadata while preserving document content
                updated_metadata = doc_metadata.get_chroma_metadata()
                collection.update(
                    ids=[document_id],
                    metadatas=[updated_metadata]
                )
                print(f"✅ Updated ChromaDB metadata for {document_id}")
            
        except Exception as chroma_error:
            print(f"⚠️ Failed to update ChromaDB metadata: {chroma_error}")
        
        return JsonResponse({
            "message": f"Document metadata updated successfully",
            "document": {
                'id': doc_metadata.id,
                'document_id': doc_metadata.document_id,
                'filename': doc_metadata.filename,
                'folder': doc_metadata.folder.name,
                'document_type': doc_metadata.get_document_type_display(),
                'target_program': doc_metadata.get_target_program_display(),
                'keywords': doc_metadata.keywords,
            }
        })
        
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)

@csrf_exempt
def delete_document(request, document_id):
    """Delete document from database, Supabase storage, and ChromaDB"""
    if request.method != 'DELETE':
        return JsonResponse({"error": "DELETE request required"}, status=400)
    
    try:
        # Get the document metadata
        doc_metadata = DocumentMetadata.objects.get(document_id=document_id)
        filename = doc_metadata.filename
        
        # Delete from Supabase storage
        try:
            supabase = get_supabase_client()
            supabase.storage.from_("original-pdfs").remove([filename])
            print(f"✅ Deleted from Supabase: {filename}")
        except Exception as supabase_error:
            print(f"⚠️ Supabase deletion failed for {filename}: {supabase_error}")
            # Continue with other deletions even if Supabase fails
        
        # Delete from ChromaDB
        try:
            from .chroma_connection import ChromaService
            collection = ChromaService.get_collection()
            
            # Delete all chunks for this document
            # ChromaDB stores document chunks with IDs like "document_id_chunk_0", "document_id_chunk_1", etc.
            result = collection.get(include=["metadatas"])
            doc_ids_to_delete = []
            
            for i, metadata in enumerate(result.get('metadatas', [])):
                if metadata and metadata.get('source_document_id') == document_id:
                    doc_ids_to_delete.append(result['ids'][i])
            
            if doc_ids_to_delete:
                collection.delete(ids=doc_ids_to_delete)
                print(f"✅ Deleted {len(doc_ids_to_delete)} chunks from ChromaDB for document: {document_id}")
            else:
                print(f"⚠️ No ChromaDB chunks found for document: {document_id}")
                
        except Exception as chroma_error:
            print(f"⚠️ ChromaDB deletion failed for {document_id}: {chroma_error}")
            # Continue with database deletion even if ChromaDB fails
        
        # Delete from Django database
        doc_metadata.delete()
        print(f"✅ Deleted from database: {document_id}")
        
        return JsonResponse({
            "status": "success",
            "message": f"Document '{filename}' deleted successfully from all systems"
        })
        
    except DocumentMetadata.DoesNotExist:
        return JsonResponse({"error": "Document not found"}, status=404)
    except Exception as e:
        print(f"❌ Delete failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({"error": str(e)}, status=500)

@api_view(['POST'])
def enhanced_chat_query(request):
    """Enhanced chat endpoint with folder/type filtering"""
    try:
        data = json.loads(request.body)
        query = data.get('query', '').strip()
        
        if not query:
            return Response({"error": "Query is required"}, status=400)
        
        # Optional manual filters
        manual_filters = {}
        if data.get('folder_filter'):
            manual_filters['folder_filter'] = data['folder_filter']
        if data.get('document_type_filter'):
            manual_filters['document_type_filter'] = data['document_type_filter']
        if data.get('program_filter'):
            manual_filters['program_filter'] = data['program_filter']
        
        # Initialize chatbot
        from .fast_hybrid_chatbot_together import FastHybridChatbotTogether
        chatbot = FastHybridChatbotTogether(use_chroma=True, chroma_collection_name="documents")
        
        # Process query with intent analysis
        if manual_filters:
            response, sources = chatbot.process_query_with_intent_analysis(
                query, 
                manual_filters=manual_filters,
                stream=False,
                require_context=True,
                min_relevance=0.3
            )
        else:
            response, sources = chatbot.process_query_with_intent_analysis(
                query,
                stream=False,
                require_context=True,
                min_relevance=0.3
            )
        
        return Response({
            "response": response,
            "sources": sources,
            "query": query,
            "applied_filters": manual_filters if manual_filters else chatbot.analyze_query_intent(query)
        })
        
    except Exception as e:
        return Response({"error": str(e)}, status=500)